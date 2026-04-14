import os 
from pathlib import Path
from typing import Union

import pandas as pd
import PyPDF2
from docx import Document

from src.domain.models import Document
from src.domain.exceptions import DocumentLoadError

# This defines the interface for all loaders
class BaseLoader:
    def load(self, file_path: Path) -> Document:
        raise NotImplementedError

# Extract from PDFs
class PDFLoader(BaseLoader):
    def load(self, file_path: Path) -> Document:
        try:
            text = ""
            with open(file_path, 'rb') as f:
                reader = PyPDF2.PdfReader(f)
                for page in reader.pages:
                    text += page.extract_text() + "\n"

            return Document(
                source_name=file_path.name,
                source_type="pdf",
                content=text.strip()
            )                
        except Exception as e:
            raise DocumentLoaderError(f"Failed to load PDF: {file_path}: {e}")    

# Extract from DOCX files
class DOCXLoader(BaseLoader):
    def load(self, file_path: Path) -> Document:
        try:
            doc = Document(file_path)
            text = "\n".join([para.text for para in doc.paragraphs])

            return Document(
                source_name=file_path.name,
                source_type="docx",
                content=text.strip()
            )            
        except Exception as e: 
            raise DocumentLoadError(f"Failed to load DOCX: {file_path}: {e}")    

# This CSV loader returns multiple Documents(csv files contain many rows, each row coulb be a separate document)
class CSVLoader(BaseLoader):
    def load(self, file_path: Path) -> list[Document]:
        try:
            df = pd.read_csv(file_path)
            document = []

            # auto-detect text columns
            text_cols = [c for c in df.columns if any(
                kw in c.lower() for kw in ["text", "content", "description", "info"]
            )]
            if not text_cols:
                text_cols = df.columns.tolist()

            for idx, row in df.iterrows():
                content = " ".join([str(row[col]) for col in text_cols if pd.notna(row[col])])
                metadata = {k: str(v) for k, v in row.items() if k not in text_cols}

                documents.append(Document(
                    source_name=f"{file_path.name}#{idx}",
                    source_type="csv",
                    content=content,
                    location=metadata.get('location'),
                    category=metadata.get('category')
                ))    

            return documents
        except Exception as e:
            raise DocumentLoadError(f"Failed to load CSV: {file_path}: {e}")    

# loader for Simple text files 
class TXTLoader(BaseLoader):
    def load(self, file_path: Path) -> Document:
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                text = f.read()

            return Document(
                source_name=file_path.name,
                source_type="txt",
                content=text.strip()
            )    
        except Exception as e:
            raise DocumentLoadError(f"Failed to load TXT: {file_path}: {e}")  

# Routes files to the correct loader based on file extension
class DocumentLoader:
    LOADERS = {
        '.pdf': PDFLoader,
        '.docx': DOCXLoader,
        '.csv': CSVLoader,
        '.txt': TXTLoader,
    } 

    def load(self, file_path: Union[str, Path]) -> Union[Document, list[Document]]:
        path = Path(file_path)
        ext = path.suffix.lower()

        if ext not in self.LOADERS:
            raise DocumentLoadError(f"Unsupported file type: {ext}")

        loader = self.LOADERS[ext]()
        result = loader.load(path)

        if isinstance(result, Document):
            return [result]
        return result

    def load_directory(self, dir_path: Union[str, Path]) -> list[Document]:
        documents = []
        path = Path(dir_path)

        for ext in self.LOADERS.keys():
            for file_path in path.glob(f"*{ext}"):
                try:
                    docs = self.load(file_path)
                    documents.extend(docs)
                    print(f" Loaded {file_path}")
                except DocumentLoadError as e:
                    print(f" Failed {file_path}: {e}")

        return documents                                         
